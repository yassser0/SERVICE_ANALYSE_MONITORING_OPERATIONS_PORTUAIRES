import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def main():
    doc = Document()

    # Title Page
    title = doc.add_heading('Port IoT - Smart Port Operations Center', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Detailed Project Architecture & Deep Learning Integration Report\n')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Section 1: Project Overview
    add_heading(doc, '1. Project Overview', 1)
    doc.add_paragraph(
        "The Port IoT Data Service API is a real-time monitoring platform for smart port operations. "
        "It acts as the central nervous system for port activities, tracking vessels, containers, and IoT sensors "
        "deployed across the port zones."
    )
    doc.add_paragraph(
        "The architecture consists of a highly scalable data pipeline utilizing Apache Kafka for message streaming, "
        "FastAPI for the backend REST API, and PostgreSQL for structured data storage. The frontend provides a "
        "glassmorphism-styled dashboard for real-time visualization of KPIs, anomalies, and zone health."
    )

    # Section 2: Architecture & Technologies
    add_heading(doc, '2. Architecture & Technologies', 1)
    doc.add_paragraph("The technology stack includes:")
    tech_list = [
        "Backend: FastAPI (Python 3.11, async architecture)",
        "Message Broker: Apache Kafka & Zookeeper (real-time data ingestion)",
        "Database: PostgreSQL with PostGIS (Time-series & relational data)",
        "Data Lake: Local file system structured in Bronze, Silver, and Gold tiers (HDFS ready)",
        "Frontend: HTML5, CSS3, Vanilla JS (Auto-refreshing dashboard)",
        "Infrastructure: Docker & Docker Compose"
    ]
    for item in tech_list:
        doc.add_paragraph(item, style='List Bullet')

    # Section 3: Data Pipeline
    add_heading(doc, '3. Data Pipeline & Flow', 1)
    doc.add_paragraph("1. Ingestion: A Simulator generates realistic IoT data (GPS, temperature, battery, speed) for various sensors and containers, pushing it to Kafka.")
    doc.add_paragraph("2. Processing: The FastAPI backend runs a background Kafka Consumer that pulls messages, validates them against Pydantic schemas, and evaluates business rules.")
    doc.add_paragraph("3. Storage: Validated data is stored in PostgreSQL tables (sensor_events, anomalies, containers, vessels). Aggregated KPIs are synced to the Data Lake (Gold tier).")
    doc.add_paragraph("4. Visualization: The dashboard polls the backend via REST endpoints (/statistics, /kpis, /anomalies) every 5 seconds to update the UI.")

    # Section 4: Anomaly Detection (Rule-Based)
    add_heading(doc, '4. Rule-Based Anomaly Detection', 1)
    doc.add_paragraph("The system currently implements 6 strict rules for anomaly detection:")
    anomalies = [
        "HIGH_TEMPERATURE: Temperature exceeds 85°C.",
        "INVALID_GPS: Coordinates fall outside the predefined port boundaries.",
        "EQUIPMENT_ERROR: Equipment reports an error or offline status.",
        "LOW_BATTERY: Sensor battery drops below 15%.",
        "SPEED_VIOLATION: Vehicle speed exceeds 25 km/h in restricted zones (e.g., Yard, Warehouse).",
        "UNEXPECTED_MOVEMENT: Container moves more than 200 meters unexpectedly (calculated via Haversine distance)."
    ]
    for a in anomalies:
        doc.add_paragraph(a, style='List Bullet')

    # Section 5: Deep Learning Integration (Advanced)
    add_heading(doc, '5. Deep Learning Integration (Future / Implementation)', 1)
    doc.add_paragraph(
        "To move beyond static rule-based alerts, the system architecture supports the integration of Deep Learning models "
        "for multivariate anomaly detection and predictive maintenance."
    )
    add_heading(doc, '5.1 Proposed Deep Learning Architecture', 2)
    doc.add_paragraph(
        "Model Type: Autoencoder (Neural Network)\n"
        "Framework: PyTorch or TensorFlow/Keras\n\n"
        "How it works: An Autoencoder is trained on historical 'normal' sensor data (features like temperature, speed, humidity, and battery level). "
        "The model learns to compress and reconstruct this normal data. When an abnormal event occurs (e.g., a subtle combination of slightly high temperature "
        "and slightly high speed that rules would miss), the model fails to reconstruct it accurately. The high 'reconstruction error' triggers a DL_ANOMALY."
    )
    add_heading(doc, '5.2 Training Workflow', 2)
    doc.add_paragraph("1. Data Extraction: Read historical sensor_events from PostgreSQL or the Data Lake.")
    doc.add_paragraph("2. Preprocessing: Normalize numerical features (MinMax Scaling) and one-hot encode categorical features (zones).")
    doc.add_paragraph("3. Training: Train the Autoencoder until the loss stabilizes.")
    doc.add_paragraph("4. Deployment: Load the trained model into the Kafka Consumer pipeline to score real-time events.")

    # Section 6: Dashboard & UI
    add_heading(doc, '6. Dashboard Overview', 1)
    doc.add_paragraph(
        "The Operations Center Dashboard provides a comprehensive view of port activities:"
    )
    dashboard_features = [
        "Live KPIs: Total Containers, Open Anomalies, Active Sensors, Vessels Docked, Flagged Containers.",
        "Zone Health Map: 5x2 grid showing traffic-light status for port zones based on 1-hour event windows.",
        "Live Anomaly Feed: Scrolling table of recent unresolved alerts.",
        "Vessels & Sensor Health: Tables showing docked ships and battery/signal status of active sensors."
    ]
    for f in dashboard_features:
        doc.add_paragraph(f, style='List Bullet')

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'Port_IoT_Report.docx')
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == "__main__":
    main()
